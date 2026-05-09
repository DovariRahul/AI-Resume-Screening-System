"""
Resume Parser Module
Extracts text from PDF and DOCX resume files.
Supports multiple PDF extraction backends for robustness.
"""

import os
import PyPDF2
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file using pdfplumber (primary)
    and PyPDF2 (fallback).
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text as a string
    """
    text = ""
    
    # Primary: Use pdfplumber for better text extraction
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"[pdfplumber] Error: {e}. Falling back to PyPDF2...")
        text = ""
    
    # Fallback: Use PyPDF2 if pdfplumber fails or extracts nothing
    if not text.strip():
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"[PyPDF2] Error extracting text: {e}")
    
    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text as a string
    """
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += row_text + "\n"
    except Exception as e:
        print(f"[python-docx] Error extracting text: {e}")
    
    return text.strip()


def parse_resume(file_path: str) -> str:
    """
    Parse a resume file and extract text based on file extension.
    
    Args:
        file_path: Path to the resume file (PDF or DOCX)
        
    Returns:
        Extracted text as a string
        
    Raises:
        ValueError: If the file format is not supported
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Use PDF or DOCX.")
